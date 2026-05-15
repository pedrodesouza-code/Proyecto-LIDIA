from __future__ import annotations

import re
from argparse import ArgumentParser
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
DEFAULT_SOURCE = ROOT / "docs" / "ENTREGA_EC3_IMPLEMENTACION.md"
DEFAULT_OUT = ROOT / "docs" / "Entrega_EC3_Implementacion.docx"


ACCENT = "1F4D78"
ACCENT_2 = "2E74B5"
MUTED = "5B677A"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "E8EEF5"
BORDER = "C9D3E1"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(clean_inline(text))
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("**", "")
    return text.strip()


def add_runs_with_inline_code(paragraph, text: str, *, bold: bool = False) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(31, 77, 120)
        else:
            run = paragraph.add_run(part.replace("**", ""))
            run.bold = bold


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in [
        ("Title", 24, ACCENT, 0, 10),
        ("Subtitle", 12, MUTED, 0, 12),
        ("Heading 1", 15, ACCENT, 14, 6),
        ("Heading 2", 12.5, ACCENT_2, 10, 4),
        ("Heading 3", 11.5, ACCENT, 8, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name.startswith("Heading") or style_name == "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = "SINIA-UY | Entrega EC3"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.size = Pt(8.5)
        p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.text = "Proyecto de Ingenieria de Datos - UTEC - 2026"
        fp.runs[0].font.size = Pt(8.5)
        fp.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(doc: Document, source: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run("SINIA-UY")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Entrega EC3 - Implementacion funcional")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Proyecto de Ingenieria de Datos | UTEC | 2026")

    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(meta, 7600)
    set_cell_margins(meta)
    set_table_borders(meta)
    rows = [
        ("Objetivo", "Demostrar funcionamiento con datos reales, ETL modular, SQL/NoSQL, CDC, testing, seguridad, dashboard, despliegue hibrido y rendimiento."),
        ("Estado", "Cumplido con evidencia versionada."),
        ("Fecha de cierre operativo", "2026-05-15"),
        ("Rama Git", "codex-flujo-local-utec"),
        ("Documento fuente", str(source.relative_to(ROOT)).replace("\\", "/")),
    ]
    for i, (label, value) in enumerate(rows):
        set_cell_shading(meta.rows[i].cells[0], LIGHT_FILL)
        set_cell_text(meta.rows[i].cells[0], label, bold=True)
        set_cell_text(meta.rows[i].cells[1], value)

    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    set_table_width(callout, 8600)
    set_cell_margins(callout, top=140, bottom=140, start=180, end=180)
    set_table_borders(callout)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_text(
        cell,
        "Resumen: EC3 queda cubierto porque el sistema ejecuta ETL con datos reales, carga PostgreSQL y MongoDB, aplica CDC e idempotencia, registra tests cuantitativos, muestra resultados en Streamlit, mide rendimiento y evidencia despliegue hibrido con UTEC.",
        bold=False,
    )

    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows = []
    for raw in table_lines:
        cells = [c.strip() for c in raw.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            continue
        rows.append(cells)
    return rows, i


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table)
    set_cell_margins(table)
    set_table_borders(table)

    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            text = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(cell, text, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor.from_string(ACCENT)
                        run.font.bold = True
    doc.add_paragraph()


def build_doc(source: Path, out: Path) -> None:
    md = source.read_text(encoding="utf-8")
    lines = md.splitlines()

    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)
    add_cover(doc, source)

    i = 0
    skip_first_title = True
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            if skip_first_title:
                skip_first_title = False
            else:
                doc.add_heading(clean_inline(stripped[2:]), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(clean_inline(stripped[3:]), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(clean_inline(stripped[4:]), level=2)
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(ACCENT)
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_inline_code(p, stripped[2:])
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs_with_inline_code(p, stripped)
        i += 1

    # Keep final page compact with a short evidence checklist.
    doc.add_heading("Checklist de entrega", level=1)
    for item in [
        "Documento EC3 generado en Word.",
        "Evidencia SQL/NoSQL versionada.",
        "Tests cuantitativos registrados.",
        "Dashboard verificado.",
        "Docker Compose validado.",
        "Credenciales reales fuera del repositorio.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.save(out)


if __name__ == "__main__":
    parser = ArgumentParser(description="Genera documento Word EC3 desde Markdown.")
    parser.add_argument("--source", default=str(DEFAULT_SOURCE), help="Markdown fuente.")
    parser.add_argument("--out", default=str(DEFAULT_OUT), help="DOCX destino.")
    args = parser.parse_args()
    source = Path(args.source)
    out = Path(args.out)
    if not source.is_absolute():
        source = ROOT / source
    if not out.is_absolute():
        out = ROOT / out
    build_doc(source, out)
    print(out)
